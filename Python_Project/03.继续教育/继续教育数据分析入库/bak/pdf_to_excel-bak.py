import pdfplumber
import pandas as pd
import time
from docx import Document

""" 继续教育数据治理导出Excel """

# 元数据格式
metadata = {
    'projectId': [], 'project_num': [], 'project_name': [], 'company': [], 'leading_name': [],
    'tel': [], 'startdate': [], 'enddate': [], 'date_num': [],
    'address': [], 'credit': [], 'crowd': [], 'person_num': [], 'remark': []
}


# 读取pdf文件，保存为pdf实例
def getPPTData():
    pdf = pdfplumber.open('../2024-0902.pdf')
    for i in range(0, 15):  # 循环PDF读取每页
        first_page = pdf.pages[i]
        table = first_page.extract_table()

        table.pop(0)  # 删除每页第一行的头文件
        for i in table:  # 构造数据
            metadata['projectId'].append('null')
            metadata['project_num'].append('%s%s' % (i[0].split('\n')[0], i[0].split('\n')[1]))
            metadata['project_name'].append(i[1])
            metadata['company'].append(i[2])
            metadata['leading_name'].append(i[3])
            metadata['tel'].append(i[4])
            metadata['address'].append(i[6])
            metadata['credit'].append(i[7].split('分')[0])
            metadata['crowd'].append(i[8])

            if i[9].split('/')[0]:
                metadata['person_num'].append(i[9].split('/')[0])
            else:
                metadata['person_num'].append(0)

            if i[10]:
                metadata['remark'].append(i[10])
            else:
                metadata['remark'].append("   ")

            metadata['startdate'].append(i[5][0:10])
            metadata['enddate'].append(i[5][11:21])
            metadata['date_num'].append(i[5][21::])


class GetMetaDate:
    # 元数据格式
    metadata = {
        'projectId': [], 'project_num': [], 'project_name': [], 'company': [], 'leading_name': [],
        'tel': [], 'startdate': [], 'enddate': [], 'date_num': [],
        'address': [], 'credit': [], 'crowd': [], 'person_num': [], 'remark': []
    }

    def read_tables_from_docx(self, world_file_path):
        """
            加载Word文档,序列化每个数据表格为列表数据,并且将序列化后的列表加入到一个大列表返回
        :return: LIST
        """
        # doc = Document('./20240902.docx')
        doc = Document(world_file_path)
        data_list = []

        for table in doc.tables:  # 遍历文档中的所有表格

            row_number = len(table.rows)  # 行数
            cells_number = len(table.rows[0].cells)  # 列数

            for i in range(1, row_number):  # # 遍历表格中的所有行,并过滤第一行表头数据
                row = table.rows[i]
                row_data = []  # 没行列表的数据

                for cell in row.cells:  # 遍历行中的所有单元格
                    data = cell.text.strip()  # 读取单元格的文本,若是字段为空赋于NULL值
                    if data:
                        row_data.append(data)
                    if not data:
                        row_data.append('NULL')
                # print(row_data)  # 打印每行数据
                data_list.append(row_data)
            # print('-' * 50)
        return data_list

    def formatting_data(self, world_file_path):
        """
            格式化元数据
        :return:
        """
        data_list = self.read_tables_from_docx(world_file_path)
        # for i in data_list[0:2]:
        for i in data_list:
            self.metadata['projectId'].append('null')

            project_num_list = i[0].split(' ')
            project_num = ''
            for j in project_num_list:
                project_num += j
            self.metadata['project_num'].append(project_num)

            self.metadata['project_name'].append(i[1].replace('\n', ''))
            self.metadata['company'].append(i[2].replace('\n', ''))
            self.metadata['leading_name'].append(i[3].replace('\n', ''))
            self.metadata['tel'].append(i[4].replace('\n', ''))
            self.metadata['address'].append(i[6].replace('\n', ''))
            self.metadata['credit'].append(i[7].split('分')[0])
            self.metadata['crowd'].append(i[8].replace('\n', ''))

            if i[9].split('/')[0]:
                self.metadata['person_num'].append(i[9].split('/')[0].replace('NULL', '0'))
            else:
                self.metadata['person_num'].append(0)

            if i[10]:
                self.metadata['remark'].append(i[10])
            else:
                self.metadata['remark'].append("   ")

            self.metadata['startdate'].append(i[5][0:10])
            self.metadata['enddate'].append(i[5].split('-')[1].split('\n')[1])
            self.metadata['date_num'].append(i[5].split('-')[1].split('\n')[2].split('天')[0].replace(' ', ''))

        return self.metadata


# 写入Excel
def to_excel(metadata):
    file_name = '非基地_%s.xlsx' % (time.strftime('%Y-%m-%d'))

    confirm_execl_data = {
        'projectId': metadata['projectId'],
        'project_num': metadata['project_num'],
        'project_name': metadata['project_name'],
        'company': metadata['company'],
        'leading_name': metadata['leading_name'],
        'tel': metadata['tel'],
        'startdate': metadata['startdate'],
        'enddate': metadata['enddate'],
        'date_num': metadata['date_num'],
        'address': metadata['address'],
        'credit': metadata['credit'],
        'crowd': metadata['crowd'],
        'person_num': metadata['person_num'],
        'remark': metadata['remark']
    }

    writer = pd.ExcelWriter(file_name)

    df_confirm = pd.DataFrame(confirm_execl_data)
    df_confirm.to_excel(excel_writer=writer, sheet_name='继续教育')

    # writer.save()
    # FutureWarning: save is not part of the public API, usage can give unexpected results and will be removed in a future version
    # 在新的版本中 save 方法已私有化 直接调用close 即可

    writer.close()


def insertDataSQL():
    # excel_file = './非基地_2024-09-02.xlsx'
    excel_file = '非基地_%s.xlsx' % (time.strftime('%Y-%m-%d'))
    sql_file_name = '非基地_INSERT_%s.sql' % (time.strftime('%Y-%m-%d'))
    df = pd.read_excel(excel_file)

    # 获取行数
    now_rows = df.shape[0]
    # print(now_rows)

    # 将空字符给赋指定默认值NULL
    df.fillna("null", inplace=True)

    with open(sql_file_name, mode='a', encoding='utf-8') as f:
        sql = "insert into edu_record (projectId, project_num, project_name, company,leading_name, tel, startdate, enddate, date_num, address, credit, crowd, person_num, remark ) values"
        f.write(sql)

        for i in range(now_rows):
            # 去除索引列
            data = list(df.loc[i].values)
            data.pop(0)

            data = str(tuple(data)) + ','
            # print(data)
            f.write(data)


if __name__ == '__main__':
    # 由PPT文件获取元数据
    getPPTData()
    print(metadata)
    # print(metadata.get('project_name'))

    # 由word获取元数据
    # a = GetMetaDate()
    # metadata = a.formatting_data(world_file_path='./2024-09-02.docx')

    # 导入EXCEL
    to_excel(metadata=metadata)

    # 读取EXCEL数据转化SQL
    # insertDataSQL()
